import os
import docx
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
from typing import List, Tuple, Dict
import logging
from tqdm import tqdm
import yaml
from dotenv import load_dotenv
import multiprocessing
from functools import partial
import psutil
from sqlalchemy import create_engine, Column, Integer, String, Float, ForeignKey
from sqlalchemy.orm import sessionmaker, relationship, declarative_base
import json
import argparse
import PyPDF2
from chunking import sentence_based_chunking

# Load environment variables
load_dotenv()

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# SQLAlchemy setup
Base = declarative_base()

class Document(Base):
    __tablename__ = 'documents'
    id = Column(Integer, primary_key=True)
    filename = Column(String)
    chunks = relationship("Chunk", back_populates="document")

class Chunk(Base):
    __tablename__ = 'chunks'
    id = Column(Integer, primary_key=True)
    document_id = Column(Integer, ForeignKey('documents.id'))
    content = Column(String)
    embedding = Column(String)  # Store as JSON string
    document = relationship("Document", back_populates="chunks")

def load_config(config_path: str) -> dict:
    with open(config_path, "r") as f:
        return yaml.safe_load(f)

def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        logging.error(f"Error extracting text from {file_path}: {str(e)}")
        return ""

def extract_text_from_pdf(file_path: str) -> str:
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            return "\n".join([page.extract_text() for page in reader.pages])
    except Exception as e:
        logging.error(f"Error extracting text from {file_path}: {str(e)}")
        return ""

def extract_text_from_txt(file_path: str) -> str:
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        logging.error(f"Error extracting text from {file_path}: {str(e)}")
        return ""

def clean_text(text: str) -> str:
    return text.replace('\n', ' ').replace('\r', '').strip()

def process_document(file_name: str, folder_path: str, config: dict) -> Dict[str, List[str]]:
    file_path = os.path.join(folder_path, file_name)
    if file_name.endswith(".docx"):
        text = extract_text_from_docx(file_path)
    elif file_name.endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
    elif file_name.endswith(".txt"):
        text = extract_text_from_txt(file_path)
    else:
        logging.warning(f"Unsupported file type: {file_name}")
        return {}

    if text:
        cleaned_text = clean_text(text)
        chunks = sentence_based_chunking(cleaned_text, config['chunk_size'], config['overlap'])
        return {file_name: chunks}
    return {}

def process_documents(folder_path: str, config: dict) -> Dict[str, List[str]]:
    cpu_count = psutil.cpu_count(logical=False)
    with multiprocessing.Pool(processes=cpu_count) as pool:
        results = list(tqdm(
            pool.imap(partial(process_document, folder_path=folder_path, config=config), os.listdir(folder_path)),
            total=len(os.listdir(folder_path)),
            desc="Processing documents"
        ))
    return {k: v for result in results for k, v in result.items()}

def create_embeddings(documents: Dict[str, List[str]], model_name: str) -> Dict[str, List[np.ndarray]]:
    try:
        model = SentenceTransformer(model_name)
    except Exception as e:
        print(f"Error loading model: {str(e)}")
        print("Please ensure you have an internet connection for the initial model download.")
        sys.exit(1)
    embeddings = {}
    for filename, chunks in tqdm(documents.items(), desc="Creating embeddings"):
        embeddings[filename] = model.encode(chunks, show_progress_bar=True)
    return embeddings

def create_faiss_index(embeddings: Dict[str, List[np.ndarray]]) -> faiss.IndexFlatL2:
    all_embeddings = np.concatenate(list(embeddings.values()))
    dimension = all_embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(all_embeddings)
    return index

def save_to_database(documents: Dict[str, List[str]], embeddings: Dict[str, List[np.ndarray]], db_path: str):
    engine = create_engine(f'sqlite:///{db_path}')
    Base.metadata.create_all(engine)
    Session = sessionmaker(bind=engine)
    session = Session()
    try:
        for filename, chunks in documents.items():
            doc = Document(filename=filename)
            session.add(doc)
            session.flush()  # To get the id of the newly created document
            for chunk, embedding in zip(chunks, embeddings[filename]):
                chunk_obj = Chunk(
                    document_id=doc.id,
                    content=chunk,
                    embedding=json.dumps(embedding.tolist())
                )
                session.add(chunk_obj)
        session.commit()
    except Exception as e:
        session.rollback()
        logging.error(f"Error saving to database: {str(e)}")
    finally:
        session.close()

def save_output(index: faiss.IndexFlatL2, index_path: str):
    faiss.write_index(index, index_path)

def main(config_path: str):
    config = load_config(config_path)
    logging.info("Starting RAG database creation process")
    
    documents = process_documents(config['folder_path'], config)
    
    if not documents:
        logging.error("No valid documents found. Exiting.")
        return
    
    total_chunks = sum(len(chunks) for chunks in documents.values())
    logging.info(f"Generated {total_chunks} chunks from {len(documents)} documents")
    
    embeddings = create_embeddings(documents, config['model_name'])
    index = create_faiss_index(embeddings)
    
    save_to_database(documents, embeddings, config['db_path'])
    save_output(index, config['output_index'])
    
    logging.info("RAG database created successfully!")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="RAG Database Creation Tool")
    parser.add_argument("config", help="Path to the configuration file")
    args = parser.parse_args()
    main(args.config)